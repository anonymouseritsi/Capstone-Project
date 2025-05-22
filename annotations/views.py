from rest_framework import viewsets
from .models import Image, Annotation, Procedure
from .serializers import ImageSerializer, AnnotationSerializer
from django.shortcuts import render, redirect, get_object_or_404
from .forms import PatientForm, ProcedureForm
from .models import *
from .forms import ImageUploadForm
import base64
from django.db.models.functions import TruncDate
from django.db.models import Count

from django.core.files.base import ContentFile

class ImageViewSet(viewsets.ModelViewSet):
    queryset = Image.objects.all()
    serializer_class = ImageSerializer

class AnnotationViewSet(viewsets.ModelViewSet):
    queryset = Annotation.objects.all()
    serializer_class = AnnotationSerializer

def home(request):
    patients = Patient.objects.all().order_by('-created_at')
    
    # Get province distribution data
    province_data = Patient.objects.values('province').annotate(count=Count('id')).order_by('province')
    labels = [item['province'] for item in province_data]
    counts = [item['count'] for item in province_data]
    
    context = {
        'patients': patients,
        'labels': labels,
        'counts': counts
    }
    return render(request,'home.html', context)

def base(request):
    return render(request,'base.html')

def annotate_image(request):
    return render(request, 'annotate.html')

def patients_view(request):
    all_patients = Patient.objects.all().order_by('-created_at')
    context = {
        'patient': all_patients,
    }
    return render(request, 'patients.html', context)

def patient_details(request, slug):
    patient = get_object_or_404(Patient, slug=slug)
    procedure = patient.procedures.first()
    #bagong dagdag
    total_cost = patient.total_procedure_cost()
    context = {
        'patient': patient,
        'procedure': procedure,
        #bagong dagdag
        'total_cost': total_cost,
    }
    return render(request, 'patient_details.html', context)

def patient_details_gallery(request, slug):
    patient = get_object_or_404(Patient, slug=slug)
    procedure = patient.procedures.first()
    #bagong dagdag
    total_cost = patient.total_procedure_cost()
    context = {
        'patient': patient,
        'procedure': procedure,
        #bagong dagdag
        'total_cost': total_cost,
    }
    return render(request, 'patient_details_gallery.html', context)

def patient_details_gallery_annotated(request, slug):
    patient = get_object_or_404(Patient, slug=slug)
    procedure = patient.procedures.first()
    #bagong dagdag
    total_cost = patient.total_procedure_cost()
    context = {
        'patient': patient,
        'procedure': procedure,
        #bagong dagdag
        'total_cost': total_cost,
    }
    return render(request, 'patient_details_gallery_annotated.html', context)

def edit_patient(request, slug):
    patient = get_object_or_404(Patient, slug=slug)
    if request.method == 'POST':
        form = PatientForm(request.POST, instance=patient)
        if form.is_valid():
            form.save()
            return redirect('patient_manager:patient_details', slug=patient.slug)
    else:
        form = PatientForm(instance=patient)
    return render(request, 'edit_patient.html', {'form': form})

def delete_patient(request, slug):
    patient = get_object_or_404(Patient, slug=slug)
    if request.method == 'POST':
        patient.delete()
        return redirect('patient_manager:patients')
    return render(request, 'delete_patient.html', {'patient': patient})

from django.db.models import Count
from django.db.models.functions import TruncDate
from .models import Procedure

def analytics_dashboard(request):
    # 1) Procedure Type Distribution (Doughnut)
    proc_dist_qs = (
        Procedure.objects
        .values('procedure_type')
        .annotate(count=Count('id'))
        .order_by('-count')
    )
    PROCEDURE_LABELS = dict(Procedure.PROCEDURE_CHOICES)
    proc_dist_labels = [PROCEDURE_LABELS.get(entry['procedure_type'], entry['procedure_type']) for entry in proc_dist_qs]
    proc_dist_counts = [entry['count'] for entry in proc_dist_qs]

    # 2) Procedures Over Time (Line chart)
    qs = (
        Procedure.objects
        .extra({'date_only': "date(date)"})
        .values('date_only', 'procedure_type')
        .annotate(count=Count('id'))
        .order_by('date_only')
    )
    types = [choice[0] for choice in Procedure.PROCEDURE_CHOICES]

    data_by_type = {ptype: {} for ptype in types}
    dates_set = set()

    for entry in qs:
        ptype = entry['procedure_type']
        date = entry['date_only']
        count = entry['count']
        data_by_type[ptype][date] = count
        dates_set.add(date)

    sorted_dates = sorted(dates_set)
    proc_time_datasets = []
    colors = ['#FF6384', '#36A2EB', '#FFCE56', '#4BC0C0', '#9966FF']

    for i, ptype in enumerate(types):
        counts = [data_by_type[ptype].get(date, 0) for date in sorted_dates]
        proc_time_datasets.append({
            'label': PROCEDURE_LABELS.get(ptype, ptype),
            'data': counts,
            'borderColor': colors[i],
            'fill': False,
            'tension': 0.2,
        })
    data = Patient.objects.values('province').annotate(count=Count('id')).order_by('province')
    labels = [item['province'] for item in data]
    counts = [item['count'] for item in data]


    context = {
        # Doughnut
        'proc_dist_labels': proc_dist_labels,
        'proc_dist_counts': proc_dist_counts,

        # Line chart
        'proc_time_labels': sorted_dates,
        'proc_time_datasets': proc_time_datasets,

        'labels': labels,  # JSON string
        'counts': counts,
    }

    
    
         

    return render(request, 'analytics_dashboard.html', context)


def billing_view(request):
    return render(request, 'billing.html')

def reports_view(request):
    return render(request, 'reports.html')

def logout_view(request):
    return render(request, 'logout.html')  



def add_patient(request):
    if request.method == 'POST':
        form = PatientForm(request.POST)
        if form.is_valid():
            patient = form.save()
            return redirect(patient.get_absolute_url())
    else:
        print("not valid")
        form = PatientForm()
    return render(request, 'add_patient.html', {'form': form})

def add_procedure(request, slug):
    patient = get_object_or_404(Patient, slug=slug)
    if request.method == 'POST':
        form = ProcedureForm(request.POST)
        if form.is_valid():
            # Create a new procedure instance and assign it to the patient
            procedure = form.save(commit=False)
            procedure.patient = patient
            procedure.save()
            return redirect('patient_manager:patient_details', slug=patient.slug)
    else:
        print("form not valid")
        form = ProcedureForm()
    
    return render(request, 'add_procedure.html', {'form': form, 'patient': patient})


def recent_patients(request):
    patients = Patient.objects.all().order_by('-created_at')[:10]  # Most recent 10
    return render(request, 'recent_patients.html', {'patients': patients})

def upload_image_for_patient(request, slug):
    patient = Patient.objects.get(slug=slug)
    if request.method == 'POST':
        form = ImageUploadForm(request.POST, request.FILES)
        if form.is_valid():
            image = form.save(commit=False)
            image.patient = patient
            image.save()
            return redirect('patient_manager:annotate', slug=patient.slug)  # or show confirmation
    else:
        form = ImageUploadForm()
    return render(request, 'upload_image.html', {'form': form, 'patient': patient})


def save_annotation(request, slug):
    if request.method == 'POST':
        patient = get_object_or_404(Patient, slug=slug)
        data_url = request.POST.get('imageData')

        if data_url:
            format, imgstr = data_url.split(';base64,') 
            ext = format.split('/')[-1]  
            data = ContentFile(base64.b64decode(imgstr), name=f'annotated_{patient.slug}.{ext}')

            latest_image = patient.images.last()
            if latest_image:
                latest_image.annotated_image.save(data.name, data)
                latest_image.save()

        return redirect('patient_manager:patient_details', slug=patient.slug)
from docx import Document
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
import requests
from tempfile import NamedTemporaryFile
from .models import Patient  # Ensure this import matches your app's structure
from django.utils.text import slugify

def download_patient_details(request, slug):
    patient = get_object_or_404(Patient, slug=slug)
    procedure = patient.procedures.first()

    document = Document()
    document.add_heading('Patient Details', 0)

    paragraph = document.add_paragraph()
    paragraph.add_run(f"Name: {patient.last_name}")
    paragraph.add_run(f"\nAge: {patient.age}")

    document.add_heading('Procedures', 1)
    if procedure:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"Procedure: {procedure.procedure_type}")
        paragraph.add_run(f"\nDate: {procedure.date}")
        paragraph.add_run(f"\nNotes: {procedure.notes}")
    else:
        document.add_paragraph("No procedure information available.")

    document.add_heading('Images', 1)
    for image in patient.images.all():
        paragraph = document.add_paragraph()
        paragraph.add_run(f"Image {image.id}:")

        # Original Image
        try:
            original_url = request.build_absolute_uri(image.image.url)
            img_response = requests.get(original_url)
            img_response.raise_for_status()

            with NamedTemporaryFile(suffix='.jpg', delete=False) as tmp:
                tmp.write(img_response.content)
                tmp.flush()
                document.add_paragraph("Original Image:")
                document.add_picture(tmp.name, width=None)
        except requests.RequestException as e:
            document.add_paragraph(f"(Original image could not be loaded: {e})")

        # Annotated Image (if available)
        if image.annotated_image:
            try:
                annotated_url = request.build_absolute_uri(image.annotated_image.url)
                ann_response = requests.get(annotated_url)
                ann_response.raise_for_status()

                with NamedTemporaryFile(suffix='.jpg', delete=False) as ann_tmp:
                    ann_tmp.write(ann_response.content)
                    ann_tmp.flush()
                    document.add_paragraph("Annotated Image:")
                    document.add_picture(ann_tmp.name, width=None)
            except requests.RequestException as e:
                document.add_paragraph(f"(Annotated image could not be loaded: {e})")

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    filename = f"{slugify(patient.last_name)}, {slugify(patient.first_name)} {slugify(patient.middle_name)}_details.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    document.save(response)

    return response



def annotate_image(request, slug):
    patient = get_object_or_404(Patient, slug=slug)
    return render(request, 'annotate.html', {'patient': patient})

